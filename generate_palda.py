import tkinter as tk
from tkinter import simpledialog, ttk, messagebox
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

# Time options
time_options = [
    "8:00AM", "8:30AM", "9:00AM", "9:30AM", "10:00AM", "10:30AM",
    "11:00AM", "11:30AM", "12:00PM", "12:30PM", "1:00PM", "1:30PM",
    "2:00PM", "2:30PM", "3:00PM", "3:30PM", "4:00PM", "4:30PM", "5:00PM"
]

save_path = r"C:\Users\ERP\Desktop\Jap\Reports_and_Was"

def generate_docx(slots):
    today = datetime.now().strftime("%m-%d-%y")
    filename = os.path.join(save_path, f"newpalda {today}.docx")

    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    try:
        table.columns[0].width = Inches(2)
    except:
        pass

    # Header
    header = table.rows[0]
    header_cell = header.cells[0]
    header_cell.merge(header.cells[1])
    title = header_cell.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Work Done")
    run.bold = True
    run.font.size = Pt(14)

    # Time slots
    for start, end in slots:
        row = table.add_row()
        time_cell = row.cells[0]
        task_cell = row.cells[1]

        # Time with formatting
        time_paragraph = time_cell.paragraphs[0]
        time_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = time_paragraph.add_run(f"-  {start} - {end}:")
        run.font.size = Pt(11)

        # Empty task cell
        task_paragraph = task_cell.paragraphs[0]
        task_paragraph.add_run("").font.size = Pt(11)

    try:
        if not os.path.exists(save_path):
            os.makedirs(save_path)
        doc.save(filename)
        os.startfile(filename)
        messagebox.showinfo("Success", f"Report saved to:\n{filename}")
    except Exception as e:
        messagebox.showerror("Error", str(e))


def ask_slot_count():
    root = tk.Tk()
    root.withdraw()
    count = simpledialog.askinteger("Time Slots", "How many time slots do you want?", minvalue=1, maxvalue=20)
    if count is None:
        messagebox.showinfo("Cancelled", "Operation cancelled.")
        exit()
    return count


def launch_gui(slot_count):
    root = tk.Tk()
    root.title("Select Time Ranges")

    # Track actual Combobox widgets instead of just StringVars
    dropdown_widgets = []

    tk.Label(root, text="Start Time", font=('Segoe UI', 10, 'bold')).grid(row=0, column=0, padx=10, pady=5)
    tk.Label(root, text="End Time", font=('Segoe UI', 10, 'bold')).grid(row=0, column=1, padx=10, pady=5)

    for i in range(slot_count):
        start_var = tk.StringVar()
        end_var = tk.StringVar()

        start_cb = ttk.Combobox(root, textvariable=start_var, values=time_options, state="readonly", width=10)
        end_cb = ttk.Combobox(root, textvariable=end_var, values=time_options, state="readonly", width=10)

        start_cb.grid(row=i+1, column=0, padx=10, pady=2)
        end_cb.grid(row=i+1, column=1, padx=10, pady=2)

        # Store Combobox widgets themselves
        dropdown_widgets.append((start_cb, end_cb))

    def on_submit():
        chosen_slots = []

        for start_cb, end_cb in dropdown_widgets:
            start = start_cb.get()
            end = end_cb.get()

            if start not in time_options or end not in time_options:
                messagebox.showerror("Missing", "Please select all time slots.")
                return

            if time_options.index(start) >= time_options.index(end):
                messagebox.showerror("Invalid", f"Start time must be before end time:\n{start} - {end}")
                return

            chosen_slots.append((start, end))

        root.destroy()
        generate_docx(chosen_slots)

    tk.Button(root, text="Generate Report", command=on_submit, bg="green", fg="white", width=20)\
        .grid(row=slot_count + 1, column=0, columnspan=2, pady=10)

    root.mainloop()


# Run everything
slot_count = ask_slot_count()
launch_gui(slot_count)
