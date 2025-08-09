from docx import Document #importing Docx for editing, adding new file docx
from docx.shared import Pt, Inches #importing inches from word
from docx.enum.text import WD_ALIGN_PARAGRAPH #importing the enum.text WD_ paragraph. this means adding constrainst,paragraph, left, justify or etc.
from datetime import datetime #import the time and date
import os

# Defining the Date today
today = datetime.now().strftime("%m-%d-%y")

# Defining the path where to save the document
save_path = r"C:\Users\ERP\Desktop\Jap\Reports_and_was"

# Creating Save path Directory using Our definition save_path that we created
if not os.path.exists(save_path): # if the file does not exist then Trigger our save_path
    try: #defining saving path, This is a triggering point, and also A validation
        os.makedirs(save_path)
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Create directory: {save_path}")
    except OSError as e:
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Error creating directory {save_path}: {e}")
        exit()

# Construct the full filename with the path
filename = os.path.join(save_path, f"newpalda {today}.docx")

# defin time slots
time_slots = [
    "8:00AM - 10:00AM",
    "10:00AM - 12:00PM",
    "1:00PM - 3:00PM",
    "3:00PM - 5:00PM"
]

# Create a new Word Document 
doc = Document()

# Create table with 2 columns
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# --- NEW ADDITION: Set column widths ---
# Adjust these values as needed to get the desired look
# A standard page width in Word (with default margins) is roughly 6.5 inches or 16.5 cm
# Let's make the first column narrower (e.g., 1.5 inches)
# The second column will expand to fill the rest of the available table width.
# For simplicity, we'll just set the first column's width.
# The second column will automatically take up the remaining space.

try:
    # Set the with of the first column
    table.columns[0].width = Inches(2) # Adjust this value (e.g., 1.0 to 2.0 inches)
    # The second column automatically adjusts to fill the remaining table width.
    # If you want to be more explicit, you could try to set the second column width too:
    # table.columns[1].width = Inches(5.0) # For example, if total table width is 6.5 inches

except IndexError:
    # this try-except block is just a safeguard; it shouldn't happen with 2 columns.
    print(f"[{datetime.now().strftime('%y-%m-%d %H:%M:%S')}] Error: Could not set column width. Table might not have enough columns. ")
    # --- END NEW ADDITION ---

    # Merge first row for the "Work Done" title
    header_row = table.rows[0]
    header_cell = header_row[0]
    header_cell.merge(header_row.cells[1])

    # Set "Work Done" title in the merged cell
    title_paragraph = header_cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run("Work Done")
    run.bold = True
    run.font.size = Pt(14)

    # Add time blocks as mew rows
    for time in time_slots:
        row = table.add_row()
        time_cell = row.cells[0]
        task_cell = row.cells[1]

    # Time and tasl format
        time_cell.text = f"- {time}:"
        task_cell.text = ""

    # Optional Formatting
    for cell in [time_cell, task_cell]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)

# save and open
try: 
    doc.save(filename)
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Document saved successfully to: {filename}")
    os.startfile(filename)
except Exception as e:
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Error saving or opening document: {e}")
